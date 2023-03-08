from docx import Document
from docx.shared import Inches
import ChessGame
from ChessReader5 import *
from ChessStatistics import *

def writeGameToDocument(dataBase):
    #VALUES
    #Won Drawn or lost
    won_draw_lost = numberOfGamesWonByStockfish(dataBase)
    won_by_Stoclfish_white = numberOfGamesWonByStockfishWhite(dataBase)
    won_by_Stockfish_black = numberOfGamesWonByStockfishBlack(dataBase)
    
    #Standard deviation and mean values
    mean_total = CalculateMeanValueOfMoves_Total(dataBase)
    mean_black = CalculateMeanValueOfMoves_StockfishBlack(dataBase)
    mean_white = CalculateMeanValueOfMoves_StockfishWhite(dataBase)
    mean_won = CalculateMeanValueOfMoves_StockfishWon(dataBase)
    mean_lost = CalculateMeanValueOfMoves_StockfishLost(dataBase)

    sd_total = CalculateStandardDeviation_Total(dataBase)
    sd_black = CalculateStandardDeviation_StockfishBlack(dataBase)
    sd_white = CalculateStandardDeviation_StockfishWhite(dataBase)
    sd_won = CalculateStandardDeviation_StockfishWon(dataBase)
    sd_lost = CalculateStandardDeviation_StockfishLost(dataBase)

    opening_dictionary = ExtractOpeningResultsofGamesPlayedNTimes(database, 20)
    document = Document()
    document.add_heading('Chessgame database statistics', 0)

    document.add_heading('Task', level=1)
    document.add_paragraph('''
    Lorem ipsum dolor sit amet, consectetur adipiscing elit. 
    In porta metus vel odio eleifend, non volutpat velit sollicitudin. Curabitur dapibus 
    mollis sem placerat lobortis. Praesent in accumsan erat. Cras eget massa quis lectus 
    dapibus hendrerit ut quis odio. Nullam non porta mauris, vitae pharetra mauris. Nulla
    quis massa ex. Integer elementum massa quis tortor fermentum molestie. Nulla imperdiet 
    nisl nisi, ac vehicula justo pharetra non. Cras laoreet neque et lobortis sodales.
    ''')


    document.add_heading('Tables', level=1)

    document.add_paragraph('\nTable regarding number of games won, lost and remis by stockfish')
    #TABLE 1
    table1 = document.add_table(rows=2, cols=3, style="Table Grid")
    heading_row = table1.rows[0].cells
    # add headings
    heading_row[0].text = "Won"
    heading_row[1].text = "Lost"
    heading_row[2].text = "Remis"
    # access second row's cells
    data_row = table1.rows[1].cells
    data_row[0].text = str(won_draw_lost[0])
    data_row[1].text = str(won_draw_lost[1])
    data_row[2].text = str(won_draw_lost[2])


    #TABLE 2
    document.add_paragraph('\nTable regarding number of games won by stockfish with white or black')
    table2 = document.add_table(rows=2, cols=2, style="Table Grid")
    heading_row1 = table2.rows[0].cells
    heading_row1[0].text = "Won with white"
    heading_row1[1].text = "Won with black"
    data_row1 = table2.rows[1].cells
    data_row1[0].text = str(won_by_Stoclfish_white)
    data_row1[1].text = str(won_by_Stockfish_black)
    
    #PlOTS
    document.add_heading('Plots', level=1)
    document.add_picture('GamesStillOngoingStockfish.png', width=Inches(5.25))
    document.add_picture('GamesStillOngoingStockfishWonorLost.png', width=Inches(5.25))

    #TABLE 3
    document.add_page_break()
    document.add_heading('Tables statistics', level=1)

    document.add_paragraph('\nTable regarding mean of the number of moves in a game')
    table3 = document.add_table(rows=6, cols=2, style="Table Grid")
    heading_row1 = table3.rows[0].cells
    heading_row1[0].text = "Mean of: "
    heading_row1[1].text = "Mean"
    data_row1 = table3.rows[1].cells
    data_row1[0].text = 'All the games in total'
    data_row1[1].text = str(mean_total)
    data_row2 = table3.rows[2].cells
    data_row2[0].text = 'Games where stockfish was white'
    data_row2[1].text = str(mean_white)    
    data_row3 = table3.rows[3].cells
    data_row3[0].text = 'Games where stockfish was white'
    data_row3[1].text = str(mean_black)    
    data_row4 = table3.rows[4].cells
    data_row4[0].text = 'Games where stockfish won'
    data_row4[1].text = str(mean_won)    
    data_row5 = table3.rows[5].cells
    data_row5[0].text = 'Games where stockfish lost'
    data_row5[1].text = str(mean_lost)


    #TABLE 4
    document.add_paragraph('\nTable regarding standard of the number of moves in a game')
    table4 = document.add_table(rows=6, cols=2, style="Table Grid")
    heading_row1 = table4.rows[0].cells
    heading_row1[0].text = "Standard deviation of: "
    heading_row1[1].text = "Standard deviation"
    data_row1 = table4.rows[1].cells
    data_row1[0].text = 'All the games in total'
    data_row1[1].text = str(sd_total)
    data_row2 = table4.rows[2].cells
    data_row2[0].text = 'Games where stockfish was white'
    data_row2[1].text = str(sd_white)    
    data_row3 = table4.rows[3].cells
    data_row3[0].text = 'Games where stockfish was white'
    data_row3[1].text = str(sd_black)    
    data_row4 = table4.rows[4].cells
    data_row4[0].text = 'Games where stockfish won'
    data_row4[1].text = str(sd_won)    
    data_row5 = table4.rows[5].cells
    data_row5[0].text = 'Games where stockfish lost'
    data_row5[1].text = str(sd_lost)



    #TABLE 4
    document.add_paragraph('\nTable regarding different openings played by stockfish and if they were won or lost')

    # Determine the number of rows needed for the table
    num_rows = len(opening_dictionary) + 1 # Add 1 for the header row
    # Add a table with four columns and num_rows rows
    table4 = document.add_table(rows=num_rows, cols=4, style="Table Grid")

    # Add column headers to the first row
    hdr_cells = table4.rows[0].cells
    hdr_cells[0].text = 'Opening'
    hdr_cells[1].text = 'Won'
    hdr_cells[2].text = 'Lost'
    hdr_cells[3].text = 'Remis'

    # Add rows to the table using data from the dictionary
    row_index = 1 # Start at row 1 (skip the header row)
    for key, value in opening_dictionary.items():
        row_cells = table4.rows[row_index].cells
        row_cells[0].text = key
        row_cells[1].text = str(value[0])
        row_cells[2].text = str(value[1])
        row_cells[3].text = str(value[2])
        row_index += 1

    # Save the document
    #document.add_page_break()
    document.save('StatDocument.docx')

database = createDataBase('/Users/erikwahlstrom/Performance_Engineering/chessassignment/Stockfish_15_64-bit.commented.[2600].pgn', 'testfil')
writeGameToDocument(database)
