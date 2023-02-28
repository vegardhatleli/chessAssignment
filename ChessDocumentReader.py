from docx import Document
from docx.shared import Inches
import ChessGame
from ChessReader5 import *

def writeGameToDocument(game):
    
    document = Document()

    document.add_heading('Chessgame number 0', 0)

    p = document.add_paragraph('Her kan vi forklare litt om hvordan oppsettet fungerer og hva vi ønsker å ta for oss ')

    document.add_heading('Trekk, hvit : svart', level=1)
    
    for move in game.Game_GetMoves():
            document.add_paragraph(
        move, style='List Bullet'
    )

    document.add_paragraph(
        'first item in unordered list', style='List Bullet'
    )

    document.add_picture('KaplanMeier.png', width=Inches(5.25))

    document.add_page_break()

    document.save('StatDocument.docx')

games = ImportChessDataBase(
    '/Users/erikwahlstrom/Performance_Engineering/chessassignment/Stockfish_15_64-bit.commented.[2600].pgn')

writeGameToDocument(game)
