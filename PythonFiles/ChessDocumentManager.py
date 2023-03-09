from docx import Document
from docx.shared import Inches
import ChessGame
from ChessReader5 import *
from ChessStatistics import *
from docx2pdf import convert
from pdf2image import convert_from_path

def writeGameToDocument(dataBase):
    #VALUES
    #Won Drawn or lost
    won_draw_lost = numberOfGamesWonByStockfish(dataBase)
    won_by_Stoclfish_white = numberOfGamesWonByStockfishWhite(dataBase)
    won_by_Stockfish_black = numberOfGamesWonByStockfishBlack(dataBase)
    
    #Standard deviation and mean values
    mean_total = CalculateMeanValueOfMoves_Total(dataBase)
    mean_black = CalculateMeanValueOfMoves_StockfishBlack(dataBase)
    mean_white = CalculateMeanValueOfMoves_StockfishWhite(dataBase)
    mean_won = CalculateMeanValueOfMoves_StockfishWon(dataBase)
    mean_lost = CalculateMeanValueOfMoves_StockfishLost(dataBase)

    sd_total = CalculateStandardDeviation_Total(dataBase)
    sd_black = CalculateStandardDeviation_StockfishBlack(dataBase)
    sd_white = CalculateStandardDeviation_StockfishWhite(dataBase)
    sd_won = CalculateStandardDeviation_StockfishWon(dataBase)
    sd_lost = CalculateStandardDeviation_StockfishLost(dataBase)


    opening_dictionary = ExtractOpeningResultsofGamesPlayedNTimes(dataBase, 20)
    
    document = Document()
    document.add_heading('Chessgame database statistics', 0)

    document.add_heading('Task', level=1)
    document.add_paragraph('''
In this task we were given to use a file that contains 2600 chess games in a pgn format (Portable Game Notation). The objective of this assignment was to design a Python script that loads these games, perform various statistics on them, a display the results at a suitable format. This document displays some of the results that we find most useful for displaying the main functionality of the script. This means that although many of the functions have many parameters that can be changed by the user we have used parameters that showcases the functionality best. All of the data for the tables and the images used are also stored in the .zip file.
    ''')


    document.add_heading('Tables', level=1)

    document.add_paragraph('\nTable regarding number of games won, lost and remis by stockfish')
    #TABLE 1
    table1 = document.add_table(rows=2, cols=3, style="Table Grid")
    heading_row = table1.rows[0].cells
    # add headings
    heading_row[0].text = "Won"
    heading_row[1].text = "Lost"
    heading_row[2].text = "Remis"
    # access second row's cells
    data_row = table1.rows[1].cells
    data_row[0].text = str(won_draw_lost[0])
    data_row[1].text = str(won_draw_lost[1])
    data_row[2].text = str(won_draw_lost[2])


    #TABLE 2
    document.add_paragraph('\nTable regarding number of games won by stockfish with white or black')
    table2 = document.add_table(rows=2, cols=2, style="Table Grid")
    heading_row1 = table2.rows[0].cells
    heading_row1[0].text = "Won with white"
    heading_row1[1].text = "Won with black"
    data_row1 = table2.rows[1].cells
    data_row1[0].text = str(won_by_Stoclfish_white)
    data_row1[1].text = str(won_by_Stockfish_black)
    
    #PlOTS
    document.add_heading('Plots', level=1)
    document.add_picture('Images/GamesStillOngoingStockfish.png', width=Inches(5.25))
    document.add_picture('Images/GamesStillOngoingStockfishWonorLost.png', width=Inches(5.25))

    #TABLE 3
    document.add_page_break()
    document.add_heading('Tables statistics', level=1)

    document.add_paragraph('\nTable regarding mean of the number of moves in a game')
    table3 = document.add_table(rows=6, cols=2, style="Table Grid")
    heading_row1 = table3.rows[0].cells
    heading_row1[0].text = "Mean of: "
    heading_row1[1].text = "Mean"
    data_row1 = table3.rows[1].cells
    data_row1[0].text = 'All the games in total'
    data_row1[1].text = str(mean_total)
    data_row2 = table3.rows[2].cells
    data_row2[0].text = 'Games where stockfish was white'
    data_row2[1].text = str(mean_white)    
    data_row3 = table3.rows[3].cells
    data_row3[0].text = 'Games where stockfish was white'
    data_row3[1].text = str(mean_black)    
    data_row4 = table3.rows[4].cells
    data_row4[0].text = 'Games where stockfish won'
    data_row4[1].text = str(mean_won)    
    data_row5 = table3.rows[5].cells
    data_row5[0].text = 'Games where stockfish lost'
    data_row5[1].text = str(mean_lost)


    #TABLE 4
    document.add_paragraph('\nTable regarding standard of the number of moves in a game')
    table4 = document.add_table(rows=6, cols=2, style="Table Grid")
    heading_row1 = table4.rows[0].cells
    heading_row1[0].text = "Standard deviation of: "
    heading_row1[1].text = "Standard deviation"
    data_row1 = table4.rows[1].cells
    data_row1[0].text = 'All the games in total'
    data_row1[1].text = str(sd_total)
    data_row2 = table4.rows[2].cells
    data_row2[0].text = 'Games where stockfish was white'
    data_row2[1].text = str(sd_white)    
    data_row3 = table4.rows[3].cells
    data_row3[0].text = 'Games where stockfish was white'
    data_row3[1].text = str(sd_black)    
    data_row4 = table4.rows[4].cells
    data_row4[0].text = 'Games where stockfish won'
    data_row4[1].text = str(sd_won)    
    data_row5 = table4.rows[5].cells
    data_row5[0].text = 'Games where stockfish lost'
    data_row5[1].text = str(sd_lost)



    #TABLE 4
    document.add_paragraph('\nTable regarding different openings played and the reusult in each of these. In this table we look at the openings that have been played more than 20 times')

    # Determine the number of rows needed for the table
    num_rows = len(opening_dictionary) + 1 # Add 1 for the header row
    # Add a table with four columns and num_rows rows
    table4 = document.add_table(rows=num_rows, cols=4, style="Table Grid")

    # Add column headers to the first row
    hdr_cells = table4.rows[0].cells
    hdr_cells[0].text = 'Opening'
    hdr_cells[1].text = 'White won'
    hdr_cells[2].text = 'Black won'
    hdr_cells[3].text = 'Remis'

    # Add rows to the table using data from the dictionary
    row_index = 1 # Start at row 1 (skip the header row)
    for key, value in opening_dictionary.items():
        row_cells = table4.rows[row_index].cells
        row_cells[0].text = key
        row_cells[1].text = str(value[0])
        row_cells[2].text = str(value[1])
        row_cells[3].text = str(value[2])
        row_index += 1

    
    image = convert_from_path('deeptree_Budapest.pdf')
    image1 = convert_from_path('shallowtree_Sicilian.pdf')

    image[0].save('Images/deeptree_Budapest.jpg', 'JPEG')
    image1[0].save('Images/shallowtree_Sicilian.jpg', 'JPEG')

    document.add_heading('Trees', level=1)
    document.add_paragraph('\nTrees for two different openenings with two different depths.')
    document.add_picture('Images/shallowtree_Sicilian.jpg', height=Inches(8))
    document.add_picture('Images/deeptree_Budapest.jpg', width=Inches(7.25))

    # Save the document as word file
    document.save('StatDocument.docx')

    #Convert to pdf and merge to one 
    convert('StatDocument.docx')

def runTest(inputLocalPathForPNGfile):
    database = createDataBase(inputLocalPathForPNGfile, 'test')
    writeGameToDocument(database)

pathForPng = '/Users/erikwahlstrom/Performance_Engineering/chessassignment/Stockfish_15_64-bit.commented.[2600].pgn'
runTest(pathForPng)